"""Generate the portfolio write-ups as Word (.docx) and Markdown (.md).

Each write-up follows a research-paper structure (Title, Author, Abstract,
Introduction, Materials & Methods, Results, Discussion, References) with
embedded, captioned figures. Run:

    python writeups/generate_writeups.py

Outputs land in writeups/docx/ and writeups/md/.
"""

from __future__ import annotations

import os
from dataclasses import dataclass, field

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

HERE = os.path.dirname(os.path.abspath(__file__))
FIG = os.path.join(HERE, "figures")
DOCX_DIR = os.path.join(HERE, "docx")
MD_DIR = os.path.join(HERE, "md")

AUTHOR = "Andrew Fogelis"


@dataclass
class Section:
    heading: str
    paras: list[str] = field(default_factory=list)
    figures: list[tuple[str, str]] = field(default_factory=list)  # (filename, caption)


@dataclass
class Doc:
    slug: str
    title: str
    abstract: str
    sections: list[Section]
    references: list[str]
    repo_url: str = ""


def build_docx(doc: Doc) -> None:
    d = Document()
    style = d.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = d.add_heading(doc.title, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    by = d.add_paragraph()
    run = by.add_run(AUTHOR)
    run.italic = True
    if doc.repo_url:
        link = d.add_paragraph()
        r = link.add_run(doc.repo_url)
        r.italic = True
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        r.font.size = Pt(9)

    d.add_heading("Abstract", level=1)
    ab = d.add_paragraph(doc.abstract)
    ab.paragraph_format.space_after = Pt(6)

    for sec in doc.sections:
        d.add_heading(sec.heading, level=1)
        for p in sec.paras:
            d.add_paragraph(p)
        for fname, caption in sec.figures:
            path = os.path.join(FIG, fname)
            if os.path.exists(path):
                d.add_picture(path, width=Inches(6.0))
                d.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap = d.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cr = cap.add_run(caption)
                cr.italic = True
                cr.font.size = Pt(9)

    d.add_heading("References", level=1)
    for ref in doc.references:
        p = d.add_paragraph(ref)
        p.paragraph_format.space_after = Pt(4)

    os.makedirs(DOCX_DIR, exist_ok=True)
    out = os.path.join(DOCX_DIR, f"{doc.slug}.docx")
    d.save(out)
    print(f"wrote {out}")


def build_md(doc: Doc) -> None:
    lines: list[str] = [f"# {doc.title}", "", f"*{AUTHOR}*", ""]
    if doc.repo_url:
        lines += [f"Repository: <{doc.repo_url}>", ""]
    lines += ["## Abstract", "", doc.abstract, ""]
    for sec in doc.sections:
        lines += [f"## {sec.heading}", ""]
        for p in sec.paras:
            lines += [p, ""]
        for fname, caption in sec.figures:
            lines += [f"![{caption}](../figures/{fname})", "", f"*{caption}*", ""]
    lines += ["## References", ""]
    for ref in doc.references:
        lines += [f"- {ref}"]
    lines += [""]
    os.makedirs(MD_DIR, exist_ok=True)
    out = os.path.join(MD_DIR, f"{doc.slug}.md")
    with open(out, "w", encoding="utf-8") as fh:
        fh.write("\n".join(lines))
    print(f"wrote {out}")


# ---------------------------------------------------------------------------
# Shared references
# ---------------------------------------------------------------------------
FOWLER = (
    "Fowler AG, Mariantoni M, Martinis JM, Cleland AN. Surface codes: Towards practical "
    "large-scale quantum computation. Physical Review A 2012; 86:032324."
)
DENNIS = (
    "Dennis E, Kitaev A, Landahl A, Preskill J. Topological quantum memory. Journal of "
    "Mathematical Physics 2002; 43:4452-4505."
)
GOOGLE = (
    "Google Quantum AI. Suppressing quantum errors by scaling a surface code logical qubit. "
    "Nature 2023; 614:676-681."
)
GIDNEY_EKERA = (
    "Gidney C, Ekera M. How to factor 2048 bit RSA integers in 8 hours using 20 million noisy "
    "qubits. Quantum 2021; 5:433."
)
PYMATCHING = (
    "Higgott O. PyMatching: A Python package for decoding quantum codes with minimum-weight "
    "perfect matching. ACM Transactions on Quantum Computing 2022; 3(3):16."
)
STIM = "Gidney C. Stim: a fast stabilizer circuit simulator. Quantum 2021; 5:497."
DELFOSSE = (
    "Delfosse N, Nickerson NH. Almost-linear time decoding algorithm for topological codes. "
    "Quantum 2021; 5:595."
)
MAAN_PALER = (
    "Maan AS, Paler A. Testing the Accuracy of Surface Code Decoders. arXiv:2311.12503, 2023."
)
SHOR = (
    "Shor PW. Algorithms for quantum computation: discrete logarithms and factoring. "
    "Proceedings 35th Annual Symposium on Foundations of Computer Science 1994; 124-134."
)
KITAEV = "Kitaev AY. Fault-tolerant quantum computation by anyons. Annals of Physics 2003; 303:2-30."
BRAVYI = (
    "Bravyi S, Suchara M, Vargo A. Efficient algorithms for maximum likelihood decoding in the "
    "surface code. Physical Review A 2014; 90:032326."
)
TORLAI = (
    "Torlai G, Melko RG. Neural decoder for topological codes. Physical Review Letters 2017; "
    "119:030501."
)
VARSAMOPOULOS = (
    "Varsamopoulos S, Criger B, Bertels K. Decoding small surface codes with feedforward neural "
    "networks. Quantum Science and Technology 2018; 3:015004."
)


DOCS: list[Doc] = [
    # -----------------------------------------------------------------
    # 0. Portfolio overview
    # -----------------------------------------------------------------
    Doc(
        slug="00_portfolio_overview",
        title="A Quantum Error Correction Research Portfolio: From Circuit-Level Simulation to Fault-Tolerance Economics",
        repo_url="https://github.com/afogelis/qec-portfolio",
        abstract=(
            "This document introduces a seven-part software portfolio in quantum error correction "
            "(QEC) that follows a single technical arc: simulate the surface code at the circuit "
            "level, decode it with classical and machine-learning algorithms, make the resulting "
            "metrics observable, and then use the same physics to answer two questions that matter "
            "for fault tolerance - how far practical decoders are from optimal, and how many "
            "physical qubits a cryptographically relevant computation would require. Each component "
            "is an independent, installable, tested and continuously integrated Python package; the "
            "later components depend on the earlier ones. Headline results obtained across the "
            "portfolio include a circuit-level threshold near 0.6%, a quantitative demonstration "
            "that belief propagation is dominated on the surface code by matching-based decoders, "
            "a calibrated reproduction of the roughly twenty-million-qubit, eight-hour estimate for "
            "Shor's algorithm on RSA-2048, a simulation reproduction of the error-suppression "
            "scaling reported by Google in 2023 (suppression factor near 2.2 below threshold), and "
            "an exact measurement of the small but growing sub-optimality of minimum-weight perfect "
            "matching relative to maximum-likelihood decoding."
        ),
        sections=[
            Section(
                "Introduction",
                [
                    "Quantum error correction is the central engineering obstacle between today's "
                    "noisy quantum processors and large-scale fault-tolerant computation. The "
                    "surface code is the leading candidate for near-term hardware because it "
                    "requires only nearest-neighbour two-qubit gates on a two-dimensional lattice "
                    "and tolerates a comparatively high physical error rate. Understanding the "
                    "surface code in depth therefore requires fluency across several disciplines: "
                    "the physics of stabiliser codes and thresholds, the algorithmics of decoding, "
                    "the statistics of rare-event estimation, and the software engineering needed "
                    "to make all of this reproducible.",
                    "This portfolio was built to demonstrate that breadth as a coherent body of "
                    "work rather than as isolated scripts. It is organised as seven independent "
                    "repositories connected by explicit dependencies, so that each can be read on "
                    "its own while the whole forms a pipeline from first-principles simulation to "
                    "decision-ready resource estimates. The remainder of this document summarises "
                    "the arc, the engineering standards shared across the repositories, and the "
                    "principal quantitative findings.",
                ],
            ),
            Section(
                "The portfolio",
                [
                    "The foundation is surface-code-simulator, a circuit-level Monte Carlo engine "
                    "built on Stim and PyMatching that constructs surface-code memory experiments, "
                    "injects noise, extracts syndromes, decodes with minimum-weight perfect "
                    "matching, and estimates the threshold. decoder-benchmark builds on it to "
                    "compare matching against from-scratch union-find and belief-propagation "
                    "decoders on accuracy, runtime and memory. ml-qec-decoder adds machine-learning "
                    "decoders (random forest, gradient-boosted trees and a neural network) and "
                    "registers them into the same benchmark for a like-for-like comparison.",
                    "qec-dashboard turns the metric artifacts emitted by the simulation and "
                    "benchmark jobs into an operational dashboard, decoupled from the heavy compute "
                    "by a small set of JSON data contracts. fault-tolerance-economics propagates "
                    "physical assumptions through the surface-code suppression law to a "
                    "physical-qubit, runtime and cost budget for Shor's algorithm on RSA-2048. The "
                    "two capstones reproduce published research: google-surface-code-reproduction "
                    "reproduces the methodology and scaling claim of Google's 2023 Nature "
                    "experiment in simulation, and decoder-accuracy-reproduction reproduces, by "
                    "exact enumeration, the matching-versus-optimal decoder comparison of Maan and "
                    "Paler (2023).",
                ],
                figures=[
                    (
                        "06_epsilon_vs_distance.png",
                        "Figure 1. Simulated logical error per cycle versus code distance (this portfolio), shown with Google's published experimental values for context.",
                    ),
                    (
                        "01_threshold_sweep.png",
                        "Figure 2. Surface-code threshold sweep produced by the foundational simulator; the distance curves cross near a physical error rate of 0.6%.",
                    ),
                ],
            ),
            Section(
                "Shared engineering practices",
                [
                    "Every repository follows the same conventions: a src/ package layout, typed "
                    "configuration objects validated with Pydantic v2, a receive-an-object / "
                    "return-an-object interface style, a command-line entry point, a pytest suite, "
                    "and a GitHub Actions workflow that runs both a ruff lint/format check and the "
                    "tests across Python 3.10 through 3.13. Randomness is seeded so that reported "
                    "numbers are reproducible, and the figures in each repository are regenerated "
                    "from the committed example scripts.",
                    "These practices were not cosmetic. Continuous integration on a clean checkout "
                    "caught a packaging defect that local testing had masked, and the discipline of "
                    "verifying every citation surfaced and corrected an attribution error in one of "
                    "the reproduction repositories. The portfolio is intended to read as "
                    "production-quality research software, not as a notebook dump.",
                ],
            ),
            Section(
                "Key results",
                [
                    "The foundational simulator places the circuit-level threshold for uniform "
                    "depolarising noise near 0.6%, consistent with the accepted range for this "
                    "noise model, and shows the expected exponential suppression of the logical "
                    "error rate with code distance below that threshold.",
                    "The decoder benchmark confirms the literature consensus that plain belief "
                    "propagation is dominated on the surface code: in a representative run, "
                    "minimum-weight perfect matching achieved the lowest mean logical error rate, "
                    "union-find was close on accuracy at near-linear cost, and belief propagation "
                    "was worse on both accuracy and runtime. The machine-learning study reaches an "
                    "honest, calibrated conclusion - learned decoders are competitive with matching "
                    "at distance three but degrade sharply at distance five under a fixed training "
                    "budget - rather than overclaiming.",
                    "The economics model reproduces the canonical roughly twenty-million-qubit, "
                    "eight-hour estimate for factoring RSA-2048 with Shor's algorithm, and "
                    "identifies the physical error rate as the dominant cost lever because it "
                    "enters the required code distance exponentially. The reproduction capstones "
                    "recover, respectively, an error-suppression factor near 2.2 below threshold "
                    "and an exact matching sub-optimality that is unity at distance three and grows "
                    "slowly at distance five.",
                ],
            ),
            Section(
                "Discussion",
                [
                    "Taken together, the repositories demonstrate the full loop a QEC researcher "
                    "works in: building trustworthy simulations, implementing and critically "
                    "comparing decoders, communicating results, and translating physics into "
                    "strategic estimates. The deliberate scoping of the reproduction projects - "
                    "reproducing methodology and qualitative conclusions in simulation rather than "
                    "claiming to match hardware-calibrated absolute numbers - is itself a "
                    "demonstration of scientific judgement.",
                    "The principal limitation shared across the portfolio is the noise model: a "
                    "single uniform depolarising rate stands in for the rich, correlated, "
                    "device-specific noise of real hardware. Natural extensions include biased and "
                    "correlated noise, leakage, more advanced decoders such as belief propagation "
                    "with ordered-statistics post-processing or correlated matching, and "
                    "larger-scale Monte Carlo via parallel samplers. These are the directions in "
                    "which each repository's future work section points.",
                ],
            ),
        ],
        references=[FOWLER, DENNIS, GOOGLE, GIDNEY_EKERA, MAAN_PALER, STIM, PYMATCHING],
    ),
    # -----------------------------------------------------------------
    # 1. surface-code-simulator
    # -----------------------------------------------------------------
    Doc(
        slug="01_surface_code_simulator",
        title="A Transparent Circuit-Level Surface-Code Memory Simulator with Threshold Estimation",
        repo_url="https://github.com/afogelis/surface-code-simulator",
        abstract=(
            "A circuit-level Monte Carlo simulator for the surface code was implemented on top of "
            "Stim and PyMatching to study the logical performance of the code as a quantum memory. "
            "The simulator builds surface-code memory circuits, applies a single physical error "
            "rate uniformly across all circuit-level noise channels, samples detection events, "
            "decodes them with minimum-weight perfect matching derived from the circuit's detector "
            "error model, and tracks logical failures with Wilson confidence intervals. A threshold "
            "sweep over code distances three, five and seven and physical error rates between 0.5% "
            "and 1.5% located the threshold near 0.6%, below which the logical error rate was "
            "suppressed with increasing code distance. The result is consistent with the accepted "
            "range for circuit-level depolarising noise and provides a transparent, testable "
            "foundation for the rest of the portfolio."
        ),
        sections=[
            Section(
                "Introduction",
                [
                    "Fault-tolerant quantum computation depends on encoding logical information so "
                    "that physical errors can be detected and corrected faster than they "
                    "accumulate. The surface code achieves this with a two-dimensional lattice of "
                    "data and measure qubits requiring only local interactions, and it tolerates a "
                    "relatively high physical error rate, which makes it the leading code for "
                    "superconducting and neutral-atom hardware. The central figure of merit is the "
                    "threshold: the physical error rate below which increasing the code distance "
                    "reduces the logical error rate.",
                    "Many high-quality libraries exist for stabiliser simulation and decoding, but "
                    "a researcher benefits from a compact, end-to-end pipeline whose statistics are "
                    "easy to read and test. This work was undertaken to build such a pipeline and "
                    "to verify that it recovers the textbook threshold behaviour, establishing a "
                    "trustworthy base layer on which decoder comparisons, machine-learning "
                    "experiments and paper reproductions could be built.",
                ],
            ),
            Section(
                "Materials and Methods",
                [
                    "Surface-code memory circuits were generated with Stim's circuit generator, "
                    "which exposes a single physical error rate that drives every circuit-level "
                    "noise channel, including two-qubit gate depolarisation, reset and measurement "
                    "flips, and idle depolarisation. Detection events were sampled with Stim's "
                    "detector sampler. The detector error model of each circuit was converted into "
                    "a matching graph and decoded with PyMatching's implementation of minimum-weight "
                    "perfect matching; a shot was counted as a logical failure when the decoded "
                    "correction disagreed with the recorded logical observable.",
                    "Logical error rates were reported with Wilson score confidence intervals, "
                    "which behave correctly for the small failure counts encountered well below "
                    "threshold, and a per-round logical error rate was derived to compare runs with "
                    "different round counts. The threshold sweep ran code distances three, five and "
                    "seven at physical error rates of 0.5%, 0.8%, 1.0%, 1.2% and 1.5% with twenty "
                    "thousand shots per point and a fixed random seed. Configuration was expressed "
                    "with typed Pydantic models, and the pipeline was covered by unit and "
                    "end-to-end tests exercising the real Stim and PyMatching stack.",
                ],
            ),
            Section(
                "Results",
                [
                    "The threshold sweep produced the characteristic crossing of logical-error "
                    "curves for different code distances. The estimated crossing fell at a physical "
                    "error rate of approximately 0.60%. Below the crossing, larger code distances "
                    "produced lower logical error rates; above it, the ordering reversed, as "
                    "expected when the code can no longer keep pace with the physical noise.",
                    "At a fixed below-threshold physical error rate of 0.8%, the logical error rate "
                    "fell steeply with code distance, displaying the exponential suppression that "
                    "is the defining benefit of the surface code. The Wilson intervals were narrow "
                    "enough at twenty thousand shots to make the ordering of the distance curves "
                    "unambiguous around the crossing region.",
                ],
                figures=[
                    (
                        "01_threshold_sweep.png",
                        "Figure 1. Logical error rate versus physical error rate for code distances three, five and seven. The curves cross near a physical error rate of 0.6%, marking the threshold.",
                    ),
                    (
                        "01_logical_vs_distance.png",
                        "Figure 2. Logical error rate versus code distance at a fixed below-threshold physical error rate of 0.8%, showing exponential suppression.",
                    ),
                ],
            ),
            Section(
                "Discussion",
                [
                    "The measured threshold near 0.6% sits within the accepted range for "
                    "circuit-level depolarising noise decoded with matching, which validates the "
                    "simulator against established results. The clean exponential suppression with "
                    "distance confirms that the noise injection, detector sampling and decoding are "
                    "wired together correctly.",
                    "The simulator's main simplification is its use of a single uniform "
                    "depolarising rate, which omits the biased, correlated and leakage-driven noise "
                    "present in hardware; absolute thresholds for real devices therefore differ. "
                    "The single-process sweep also trades raw throughput for transparency. Natural "
                    "extensions include parallelised sampling for larger sweeps, biased-noise "
                    "channels, and alternative decoders, the last of which is addressed directly by "
                    "the companion decoder-benchmark repository.",
                ],
            ),
        ],
        references=[FOWLER, DENNIS, STIM, PYMATCHING, KITAEV],
    ),
    # -----------------------------------------------------------------
    # 2. decoder-benchmark
    # -----------------------------------------------------------------
    Doc(
        slug="02_decoder_benchmark",
        title="Benchmarking Surface-Code Decoders: Minimum-Weight Matching, Union-Find and Belief Propagation",
        repo_url="https://github.com/afogelis/decoder-benchmark",
        abstract=(
            "A benchmarking framework was developed to compare surface-code decoders on a level "
            "playing field, scoring each on accuracy, runtime and peak memory over identical "
            "batches of error syndromes. Minimum-weight perfect matching (via PyMatching) was "
            "compared against from-scratch implementations of a union-find decoder and a "
            "log-domain belief-propagation decoder. Across code distances three and five and "
            "physical error rates between 0.5% and 1.2%, matching achieved the lowest mean logical "
            "error rate, union-find was close on accuracy at near-linear time cost, and belief "
            "propagation was dominated on both accuracy and runtime. The framework reproduces the "
            "consensus of the decoder literature and exposes, by implementing the algorithms "
            "directly, why each behaves as it does."
        ),
        sections=[
            Section(
                "Introduction",
                [
                    "A decoder maps the syndrome produced by a quantum error-correcting code to a "
                    "correction. Its quality determines how close a code operates to its "
                    "theoretical threshold, and its speed determines whether decoding can keep pace "
                    "with the measurement cycle in real time. The surface code admits several "
                    "decoding strategies with different accuracy and latency trade-offs, so a fair "
                    "comparison framework is valuable both pedagogically and practically.",
                    "This work implemented union-find and belief-propagation decoders from scratch, "
                    "rather than calling compiled libraries, so that the benchmark could illuminate "
                    "the algorithmic reasons behind each decoder's performance instead of treating "
                    "it as a black box. Matching, the community-standard accuracy reference, was "
                    "included through the established PyMatching library.",
                ],
            ),
            Section(
                "Materials and Methods",
                [
                    "The detector error model emitted by each surface-code circuit was converted "
                    "into parity-check and observable matrices shared by every decoder, ensuring "
                    "that all decoders saw identical syndrome batches. Matching was performed with "
                    "PyMatching. The union-find decoder implemented the Delfosse-Nickerson "
                    "cluster-growth algorithm with spanning-forest peeling on the matching graph, "
                    "which runs in almost-linear time. The belief-propagation decoder implemented "
                    "the log-domain sum-product algorithm on the detector error model.",
                    "Each decoder was profiled for accuracy (logical error rate), runtime "
                    "(microseconds per shot) and peak memory (via allocation tracking) over the "
                    "same shots. Runs were summarised in a ranked leaderboard and an "
                    "accuracy-versus-runtime Pareto frontier. The reported configuration covered "
                    "code distances three and five at physical error rates of 0.5%, 0.8%, 1.0% and "
                    "1.2% with five thousand shots per point and a fixed seed.",
                ],
            ),
            Section(
                "Results",
                [
                    "Minimum-weight perfect matching achieved the best mean logical error rate of "
                    "the three decoders and the lowest runtime, reflecting its optimised compiled "
                    "implementation. The union-find decoder was close to matching on accuracy while "
                    "using markedly less memory, consistent with its near-linear design, but was "
                    "slower than the compiled matching routine in this pure-Python implementation. "
                    "Belief propagation was the weakest decoder, with the highest logical error "
                    "rate and the highest runtime, leaving it dominated on both axes of the Pareto "
                    "frontier.",
                    "The accuracy-versus-physical-error-rate curves at distance five reproduced the "
                    "expected ordering across the full range of physical error rates studied, with "
                    "matching and union-find defining the accuracy frontier and belief propagation "
                    "trailing.",
                ],
                figures=[
                    (
                        "02_pareto.png",
                        "Figure 1. Accuracy-versus-runtime Pareto frontier. Matching and union-find define the accuracy frontier; belief propagation is dominated on both axes.",
                    ),
                    (
                        "02_accuracy_vs_p_d5.png",
                        "Figure 2. Logical error rate versus physical error rate at code distance five for each decoder.",
                    ),
                ],
            ),
            Section(
                "Discussion",
                [
                    "That plain belief propagation is dominated on the surface code is a known "
                    "result, and reproducing it from a direct implementation clarifies the cause: "
                    "the surface-code factor graph is highly degenerate and rich in short cycles, "
                    "which prevents the message-passing iteration from converging to the correct "
                    "marginal in the way it does for the sparse, loop-poor graphs of classical "
                    "low-density parity-check codes. Belief propagation becomes competitive only "
                    "when augmented, for example with ordered-statistics post-processing.",
                    "Union-find's strong accuracy at near-linear theoretical cost makes it "
                    "attractive for real-time decoding; the runtime gap observed here reflects the "
                    "pure-Python implementation rather than the algorithm itself. Future work "
                    "includes adding belief propagation with ordered-statistics decoding and "
                    "correlated matching, and compiling the union-find inner loop. The "
                    "machine-learning decoders evaluated in the companion repository register into "
                    "this same framework for direct comparison.",
                ],
            ),
        ],
        references=[DENNIS, DELFOSSE, PYMATCHING, FOWLER, STIM],
    ),
    # -----------------------------------------------------------------
    # 3. qec-dashboard
    # -----------------------------------------------------------------
    Doc(
        slug="03_qec_dashboard",
        title="An Observability Dashboard for Surface-Code Simulation Metrics",
        repo_url="https://github.com/afogelis/qec-dashboard",
        abstract=(
            "An operational dashboard was built to make the metrics produced by surface-code "
            "simulation and decoder-benchmark jobs observable in the way an operations team would "
            "monitor a production system. Implemented with Streamlit, the dashboard consumes a "
            "small set of JSON artifacts - threshold sweeps, decoder leaderboards and syndrome "
            "statistics - through explicit data contracts, so that it remains decoupled from the "
            "heavy simulation stack. It presents physical-versus-logical error rates by code "
            "distance, code-distance effects, per-detector syndrome statistics, decoder "
            "performance and a raw run explorer, and ships with bundled sample data so that it runs "
            "with no simulations required."
        ),
        sections=[
            Section(
                "Introduction",
                [
                    "Simulation and benchmarking produce numbers; turning those numbers into "
                    "decision-ready views is a separate skill. In operational settings, dashboards "
                    "consume metrics emitted by upstream pipelines rather than recomputing them, "
                    "which keeps the presentation layer fast, cheap to deploy and robust to changes "
                    "in the compute layer. This work applied that pattern to quantum error "
                    "correction.",
                    "The goal was a dashboard that an error-correction operations team would "
                    "actually watch, surfacing the quantities that indicate whether a code is "
                    "operating below threshold and which decoder is performing best, while "
                    "remaining independent of the simulators that generate the data.",
                ],
            ),
            Section(
                "Materials and Methods",
                [
                    "The dashboard depends only on a set of JSON schemas describing three "
                    "artifacts: a threshold sweep, a decoder benchmark and a syndrome-statistics "
                    "summary. Loader functions parse these artifacts into tabular frames, and the "
                    "interface was implemented as a set of Streamlit tabs. Because the contract is "
                    "the JSON schema rather than the simulation code, the dashboard installs and "
                    "runs without the simulation dependencies.",
                    "Sample artifacts were bundled with the application so that it works out of the "
                    "box; an optional generator regenerates them from the upstream simulator and "
                    "benchmark packages. Loader behaviour was covered by unit tests.",
                ],
            ),
            Section(
                "Results",
                [
                    "The dashboard presents five views: an overview of physical-versus-logical "
                    "error rates by code distance with the threshold marked, a code-distance-effects "
                    "view of logical error rate versus distance at a chosen physical rate, "
                    "per-detector syndrome firing statistics, a decoder-performance view combining "
                    "the leaderboard and the accuracy/runtime trade-off, and a run explorer for "
                    "filtering and exporting raw records. Running on the bundled sample data, the "
                    "overview reproduces the threshold crossing and distance ordering from the "
                    "underlying simulator.",
                ],
                figures=[
                    (
                        "03_dashboard_overview.png",
                        "Figure 1. The dashboard overview tab running on bundled sample data: summary metrics and logical-versus-physical error-rate curves by code distance, with the estimated threshold marked.",
                    ),
                ],
            ),
            Section(
                "Discussion",
                [
                    "Separating presentation from computation through a small data contract is the "
                    "same discipline that underlies production observability stacks, and it pays "
                    "off here in deployability and testing speed. The dashboard is intentionally "
                    "read-only and static; it visualises completed runs rather than streaming live "
                    "results.",
                    "Future work includes streaming metrics from long-running sweeps, alerting when "
                    "an operating point drifts above threshold, and richer cross-filtering across "
                    "decoders and code distances. The clean contract makes such extensions additive "
                    "rather than invasive.",
                ],
            ),
        ],
        references=[FOWLER, STIM, PYMATCHING, GOOGLE],
    ),
    # -----------------------------------------------------------------
    # 4. ml-qec-decoder
    # -----------------------------------------------------------------
    Doc(
        slug="04_ml_qec_decoder",
        title="Machine-Learning Decoders for the Surface Code: A Regime Analysis against Minimum-Weight Perfect Matching",
        repo_url="https://github.com/afogelis/ml-qec-decoder",
        abstract=(
            "Three machine-learning decoders for the surface code - a random forest, a "
            "gradient-boosted tree ensemble and a feed-forward neural network - were implemented "
            "and compared head-to-head with minimum-weight perfect matching. Each model learns to "
            "predict the logical observable flip directly from a syndrome and plugs into the "
            "decoder-benchmark framework for a like-for-like comparison. Across code distances "
            "three and five and physical error rates between 1% and 3% with a fixed training "
            "budget, the learned decoders were competitive with matching at distance three - the "
            "neural network reached a logical error rate of 0.080 against matching's 0.064 at a "
            "physical rate of 1% - but degraded sharply at distance five. The study reaches a "
            "calibrated conclusion about when learned decoding helps rather than overclaiming that "
            "it beats matching."
        ),
        sections=[
            Section(
                "Introduction",
                [
                    "Machine-learning decoders are an active research direction for quantum error "
                    "correction, motivated by the hope that a learned model can capture noise "
                    "correlations that hand-designed decoders ignore. Early neural decoders showed "
                    "promise on small codes, and the question of when learning helps remains "
                    "practically important.",
                    "This work framed decoding as supervised classification from syndrome to "
                    "logical flip and asked a focused question: under a fixed, realistic training "
                    "budget and the same circuit-level depolarising noise used elsewhere in the "
                    "portfolio, in which regimes do learned decoders match or beat minimum-weight "
                    "perfect matching, and where do they fail?",
                ],
            ),
            Section(
                "Materials and Methods",
                [
                    "Three models were implemented behind a common base class: a random forest and "
                    "a gradient-boosted tree ensemble, and a feed-forward neural network trained "
                    "with binary cross-entropy loss, the Adam optimiser and early stopping on a "
                    "validation split. Training data were sampled from the same Stim circuits the "
                    "classical decoders see, so the comparison is apples-to-apples, and the models "
                    "were registered into the decoder-benchmark framework to be scored identically.",
                    "The reported sweep covered code distances three and five at physical error "
                    "rates of 1%, 1.5%, 2% and 3%, with twenty thousand training shots and five "
                    "thousand evaluation shots at a fixed seed. Minimum-weight perfect matching was "
                    "evaluated on the same shots as the reference.",
                ],
            ),
            Section(
                "Results",
                [
                    "At code distance three the learned decoders were competitive with matching. "
                    "The neural network achieved a logical error rate of 0.080 against matching's "
                    "0.064 at a physical error rate of 1%, and the gap closed further with "
                    "additional training data; inference was sub-microsecond per shot because a "
                    "forward pass is a few matrix multiplications. At code distance five every "
                    "learned decoder degraded markedly - for example the best learned decoder "
                    "reached 0.335 against matching's 0.080 at a physical rate of 1% - because the "
                    "syndrome space grows, logical flips become rarer, and a fixed training budget "
                    "no longer covers the input distribution.",
                    "Across all regimes matching won, but the margin and the reasons varied with "
                    "distance and physical error rate, producing a clear regime map rather than a "
                    "single verdict.",
                ],
                figures=[
                    (
                        "04_ml_vs_mwpm.png",
                        "Figure 1. Logical error rate of matching versus the best machine-learning decoder, across code distance and physical error rate. The learned decoders approach matching at distance three and low physical rates, then fall behind at distance five.",
                    ),
                ],
            ),
            Section(
                "Discussion",
                [
                    "For circuit-level depolarising noise the matching graph is an excellent model "
                    "of the error process, so a learned decoder is competing against a "
                    "near-optimal baseline; matching also exploits the known error model rather "
                    "than having to learn it from data. The honest conclusion is therefore not that "
                    "machine learning beats matching, but that it is competitive only where the "
                    "matching graph is a poor model - strongly correlated or non-graphlike noise - "
                    "or where training data are abundant relative to the code distance.",
                    "The fixed training budget is the central limitation; performance at distance "
                    "five is data-starved by construction. Future work includes scaling training "
                    "data with distance, convolutional and graph-neural architectures that exploit "
                    "lattice locality, and evaluation under correlated and leakage noise where "
                    "learned models are most likely to add value.",
                ],
            ),
        ],
        references=[TORLAI, VARSAMOPOULOS, PYMATCHING, FOWLER, STIM],
    ),
    # -----------------------------------------------------------------
    # 5. fault-tolerance-economics
    # -----------------------------------------------------------------
    Doc(
        slug="05_fault_tolerance_economics",
        title="Physical-Qubit and Runtime Economics of Shor's Algorithm on RSA-2048",
        repo_url="https://github.com/afogelis/fault-tolerance-economics",
        abstract=(
            "A transparent resource and cost model was developed to estimate the physical qubits, "
            "runtime and cost required to run Shor's algorithm against RSA-2048 under realistic "
            "error-correction overhead. Logical resource requirements were taken from the "
            "literature and propagated through the surface-code suppression law to fix the code "
            "distance, the per-patch physical-qubit footprint and the total runtime. Calibrated to "
            "the canonical estimate of Gidney and Ekera (2021), a baseline superconducting profile "
            "yielded roughly twenty-three million physical qubits at code distance twenty-nine and "
            "a runtime of about seven to eight hours. A sensitivity analysis identified the "
            "physical error rate as the dominant cost lever, because it enters the required code "
            "distance exponentially."
        ),
        sections=[
            Section(
                "Introduction",
                [
                    "The security of widely deployed public-key cryptography rests on the classical "
                    "hardness of factoring. Shor's algorithm would break RSA on a sufficiently "
                    "large fault-tolerant quantum computer, so the physical-resource cost of "
                    "running it is a question of direct strategic interest. That cost is dominated "
                    "not by the logical algorithm but by the error-correction overhead needed to "
                    "execute it reliably.",
                    "This work built a transparent model that turns explicit physical assumptions "
                    "into a physical-qubit, runtime and cost budget, with the aim of reproducing "
                    "the canonical published estimate and exposing which assumptions matter most.",
                ],
            ),
            Section(
                "Materials and Methods",
                [
                    "Logical resource requirements - the number of algorithmic logical qubits and "
                    "the Toffoli count for Shor on RSA-2048, together with a factory and routing "
                    "multiplier - were taken from Gidney and Ekera (2021). The surface-code "
                    "overhead was modelled with the standard suppression law, in which the logical "
                    "error per patch scales as roughly one tenth of the ratio of physical to "
                    "threshold error rate raised to the power of half the distance plus one; a "
                    "total error budget fixed the required code distance, and a rotated patch was "
                    "costed at twice the distance squared minus one physical qubits.",
                    "Runtime was estimated as the Toffoli count multiplied by a per-Toffoli time, "
                    "calibrated so that the baseline profile reproduced the published figure. "
                    "Hardware assumptions were captured in typed profiles, and a sensitivity sweep "
                    "varied the physical error rate, cycle time and threshold to produce optimistic, "
                    "baseline and conservative scenarios.",
                ],
            ),
            Section(
                "Results",
                [
                    "The baseline superconducting profile - a physical error rate of one in a "
                    "thousand, a one-microsecond surface-code cycle and a one-percent threshold - "
                    "yielded approximately twenty-three million physical qubits at code distance "
                    "twenty-nine, running for about seven and a half hours, in line with the "
                    "canonical roughly twenty-million-qubit, eight-hour estimate. An optimistic "
                    "profile reduced the requirement to under ten million qubits at distance "
                    "nineteen, while a conservative profile raised it above eighty million qubits "
                    "at distance fifty-five.",
                    "The sensitivity analysis showed that the physical error rate dominates the "
                    "budget: because it enters the required distance exponentially, modest "
                    "improvements in physical fidelity translate into large reductions in "
                    "physical-qubit count, far outweighing changes in cycle time or the assumed "
                    "threshold.",
                ],
                figures=[
                    (
                        "05_sensitivity.png",
                        "Figure 1. Sensitivity of the physical-qubit estimate to each modelling assumption. The physical error rate dominates because it enters the required code distance exponentially.",
                    ),
                ],
            ),
            Section(
                "Discussion",
                [
                    "Reproducing the canonical estimate from an independent, transparent model "
                    "increases confidence in both the estimate and the model, and the dominance of "
                    "the physical error rate clarifies where hardware effort yields the greatest "
                    "leverage. The exercise also illustrates translating a physics result into a "
                    "decision-ready figure with explicit, inspectable assumptions.",
                    "The model is deliberately simplified: it abstracts magic-state distillation, "
                    "routing congestion and lattice-surgery scheduling into multipliers and a "
                    "calibrated per-Toffoli time rather than modelling them in detail. Future work "
                    "includes an explicit distillation-factory model, a scheduling-aware runtime "
                    "estimate, and profiles for additional hardware modalities such as neutral "
                    "atoms and trapped ions.",
                ],
            ),
        ],
        references=[SHOR, GIDNEY_EKERA, FOWLER, KITAEV],
    ),
    # -----------------------------------------------------------------
    # 6. google reproduction
    # -----------------------------------------------------------------
    Doc(
        slug="06_google_reproduction",
        title="Reproducing the Scaling Logic of Google's 2023 Surface-Code Experiment in Simulation",
        repo_url="https://github.com/afogelis/google-surface-code-reproduction",
        abstract=(
            "The central scaling claim of Google Quantum AI's 2023 experiment - that below "
            "threshold, increasing the surface-code distance suppresses the logical error per "
            "cycle - was reproduced in simulation. Logical error per cycle was extracted by fitting "
            "the decay of logical fidelity with the number of error-correction rounds, using the "
            "portfolio's Stim and matching simulator, for code distances three, five and seven at a "
            "representative below-threshold physical error rate. The simulation reproduced the "
            "qualitative suppression with an error-suppression factor near 2.2 between successive "
            "distances. Consistent with honest scoping, device-specific absolute error rates were "
            "not reproduced, because a single uniform depolarising model was used in place of "
            "Google's calibrated per-component noise; the published values are shown only for "
            "context."
        ),
        sections=[
            Section(
                "Introduction",
                [
                    "In 2023 Google Quantum AI reported the first experimental demonstration that a "
                    "larger surface code can have a lower logical error rate than a smaller one, a "
                    "milestone for the field because it showed error suppression by scaling on real "
                    "hardware. The result is naturally summarised by a suppression factor relating "
                    "the logical error per cycle at successive code distances.",
                    "A full hardware reproduction is impossible without the device, so this work "
                    "set out to reproduce the experiment's analysis methodology and its qualitative "
                    "physics conclusion in simulation, and to be explicit about the boundary "
                    "between what simulation can and cannot reproduce.",
                ],
            ),
            Section(
                "Materials and Methods",
                [
                    "The logical error per cycle was extracted by fitting the relation that one "
                    "minus twice the failure probability after a given number of rounds equals one "
                    "minus twice the per-cycle logical error, raised to the power of the round "
                    "count. The failure probability as a function of round count was measured with "
                    "the portfolio's Stim and matching simulator. Code distances three, five and "
                    "seven were run at a representative below-threshold physical error rate of "
                    "0.4%. A weighted mean of representative published component error rates was "
                    "used to place the simplified uniform model near the experiment's operating "
                    "regime.",
                ],
            ),
            Section(
                "Results",
                [
                    "The extracted logical error per cycle fell monotonically with code distance, "
                    "from approximately 0.0035 at distance three to 0.0016 at distance five and "
                    "0.0007 at distance seven. The corresponding error-suppression factors were "
                    "approximately 2.2 between distances three and five and 2.3 between distances "
                    "five and seven, reproducing the experiment's central claim that increasing the "
                    "code distance suppresses the logical error per cycle below threshold.",
                    "The published experimental values, near three percent per cycle, were plotted "
                    "alongside the simulation for context. As expected, the absolute simulated "
                    "values differ from the hardware values, since the simulation uses uniform "
                    "depolarising noise rather than the device's calibrated, correlated noise.",
                ],
                figures=[
                    (
                        "06_fidelity_decay.png",
                        "Figure 1. Logical error versus number of error-correction rounds for code distances three, five and seven; the per-cycle logical error is fit from each decay curve.",
                    ),
                    (
                        "06_epsilon_vs_distance.png",
                        "Figure 2. Simulated logical error per cycle versus code distance, with Google's published experimental values overlaid for context. Below threshold the simulated value falls with distance.",
                    ),
                ],
            ),
            Section(
                "Discussion",
                [
                    "The simulation reproduces the qualitative result that matters - error "
                    "suppression by scaling - and recovers a suppression factor in the same range "
                    "as the experiment. The deliberate decision not to claim reproduction of the "
                    "absolute numbers is the scientifically honest position: matching "
                    "hardware-calibrated values would require the device's detailed noise model, "
                    "which is not available in this setting.",
                    "Limitations include the uniform depolarising noise model, the omission of "
                    "leakage, crosstalk and non-Markovian effects present in hardware, and the use "
                    "of matching as the sole decoder. Future work includes substituting a "
                    "calibrated component-wise noise model and studying the more advanced decoders "
                    "the experiment also considered.",
                ],
            ),
        ],
        references=[GOOGLE, FOWLER, STIM, PYMATCHING],
    ),
    # -----------------------------------------------------------------
    # 7. decoder-accuracy reproduction
    # -----------------------------------------------------------------
    Doc(
        slug="07_decoder_accuracy_reproduction",
        title="Exact Sub-Optimality of Minimum-Weight Perfect Matching versus Maximum-Likelihood Decoding",
        repo_url="https://github.com/afogelis/decoder-accuracy-reproduction",
        abstract=(
            "The methodology of Maan and Paler (2023), which compares practical decoders against an "
            "exact reference, was reproduced by enumerating every error pattern of small "
            "surface codes in the code-capacity model. Because the number of independent error "
            "mechanisms is small at low distance, the logical error rate of any decoder can be "
            "computed exactly without Monte Carlo sampling. The optimal maximum-likelihood decoder, "
            "which selects the most probable logical class per syndrome, sets a hard lower bound; "
            "minimum-weight perfect matching was evaluated against it. Matching was found to be "
            "exactly optimal at code distance three across all physical error rates studied, and to "
            "develop a small, growing sub-optimality at distance five, reaching a factor of about "
            "1.004 at a physical error rate of fifteen percent."
        ),
        sections=[
            Section(
                "Introduction",
                [
                    "Every practical decoder trades accuracy for speed, but quantifying that "
                    "trade-off requires a reference of known quality. The optimal decoder for a "
                    "given code and noise model is the maximum-likelihood decoder, which chooses "
                    "the most probable logical class consistent with the observed syndrome; its "
                    "logical error rate is a hard lower bound that no decoder can beat.",
                    "Maan and Paler (2023) compared matching and belief propagation against "
                    "exhaustive look-up tables for surface codes up to distance seven. This work "
                    "reproduced the core of that methodology by exact enumeration on small "
                    "code-capacity instances, in order to measure precisely how far matching sits "
                    "from optimal and how that gap grows with the physical error rate.",
                ],
            ),
            Section(
                "Materials and Methods",
                [
                    "The analysis used the code-capacity model: a single round of correction with "
                    "data-qubit noise and perfect stabiliser measurements. In this model a small "
                    "surface code has few enough independent error mechanisms that every error "
                    "pattern can be enumerated. For each syndrome, the total probability of each "
                    "logical class was summed over all consistent error patterns; the optimal "
                    "decoder's error rate is one minus the sum over syndromes of the maximum "
                    "per-class probability. Minimum-weight perfect matching was evaluated on the "
                    "same exact enumeration, yielding its exact logical error rate rather than a "
                    "sampled estimate. Code distances three and five were studied across physical "
                    "error rates from two to fifteen percent.",
                ],
            ),
            Section(
                "Results",
                [
                    "At code distance three the matching logical error rate equalled the optimal "
                    "bound exactly across every physical error rate studied, giving a "
                    "sub-optimality ratio of one. At code distance five matching remained on the "
                    "optimal bound at low physical error rates and developed a small sub-optimality "
                    "as the rate increased, reaching a ratio of approximately 1.004 at a physical "
                    "error rate of fifteen percent. The matching logical error rate was at or above "
                    "the optimal bound everywhere, as it must be.",
                ],
                figures=[
                    (
                        "07_ler_vs_p.png",
                        "Figure 1. Exact logical error rate of the optimal maximum-likelihood decoder and of matching versus physical error rate. At low physical error rates the curves coincide.",
                    ),
                    (
                        "07_suboptimality.png",
                        "Figure 2. Sub-optimality ratio (matching divided by optimal, at least one) versus physical error rate. Matching is exactly optimal at distance three and develops a small, growing gap at distance five.",
                    ),
                ],
            ),
            Section(
                "Discussion",
                [
                    "The results reproduce the central message of the source paper: matching is an "
                    "excellent but not strictly optimal decoder, and its sub-optimality is small "
                    "and quantifiable. The gap appears and grows where degenerate error "
                    "configurations - which matching cannot weigh against one another - become more "
                    "important, namely at higher physical error rates and larger code distances.",
                    "Exact enumeration restricts the analysis to small distances in the "
                    "code-capacity model, which omits measurement errors and multi-round dynamics. "
                    "The source paper reaches distance seven with exhaustive look-up tables; "
                    "tensor-network maximum-likelihood decoding would extend the exact comparison to "
                    "larger codes and to the circuit-level model, which is the natural direction for "
                    "future work.",
                ],
            ),
        ],
        references=[MAAN_PALER, BRAVYI, DENNIS, PYMATCHING],
    ),
]


def main() -> None:
    for doc in DOCS:
        build_docx(doc)
        build_md(doc)
    print(f"\nGenerated {len(DOCS)} write-ups in {DOCX_DIR} and {MD_DIR}")


if __name__ == "__main__":
    main()
